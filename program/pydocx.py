from docx import Document

Docx = Document()
# Docx = Document('d:/test.docx')  #打开文档
Docx.add_heading("这是一个一级标题",level=1)
Docx.add_paragraph("这是一个副级标题","Title")
A = Docx.add_paragraph("My name is aaa")
A.add_run("我学习的很快乐，啊哈哈哈哈哈，非常好 Good!!!")
Docx.add_heading("这是一个二级标题",level=2)
A = Docx.add_paragraph("这个是二级标题的内容呀")
B = A.add_run("二级标题里面的正文 继续添加！！！！！！！")
B.font.bold = True # 同时我要对这些正文进行加粗~~~~
B.font.size = (20)
Docx.add_heading("我爱学习Python以下就是python的logo呀",level=3)
# Docx.add_picture("1.png")
Docx.add_table(rows=5, cols=5)
Docx.save("file/Python.docx")